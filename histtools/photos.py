# -*- coding: utf-8 -*-
"""
Created on Fri Mar 13 08:34:31 2015

History Photo Tools

Suite of tools for working with environmental photos

@author: bbatt
"""


from arcpy import da, SpatialReference, CreateUniqueName
from PIL import Image
import os
import docx

#Define Geodatabase with attachments
#gdb = "filegdb"
#Define where to save the attachments
#photows = "photo_location"
#if not os.path.exists(photows):
#    os.mkdir(photows)
#Name of feature class with attachments
#fcname = "HistoricStructures"
#fc = os.path.join(gdb, fcname)
#Specify sql for table report
#sql = "Eligibility = 'E'"

#Define spatial reference for resource table
#WGS_1984 UTM 16 WKID
#spatref = SpatialReference(32716)
#WGS_1984 UTM 17 WKID
#spatref = SpatialReference(32617)

subtypes = {
0: "Barn", 
1: "Cemetery",
2: "Church",
3: "Commercial",
4: "Gas Station",
5: "House",
6: "Other",
7: "Outbuilding",
8: "School",
}

eligdict = {
"E": "Eligible",
"NE": "Not Eligible",
"L": "Listed",
"PE": "Recommended Eligible",
"U":"Unknown"
}

style = {
'COL' : 'Colonial Revival',
'CRF' : 'Craftsman',
'ECR' : 'Early Classical Revival',
'EVR' : 'English Vernacular Revival',
'FE' : 'Federal',
'FR' : 'Federal Revival',
'FV' : 'Folk Victorian',
'FVR' : 'French Vernacular Revival',
'GE' : 'Georgian',
'GOV' : 'Gothic Revival',
'GRV' : 'Greek Revival',
'HVE' : 'High Victorian Eclectic',
'INT' : 'International ',
'IRR' : 'Italian Renaissance Revival',
'ITA' : 'Italianate',
'MR' : 'Mediterranean Revival',
'NR' : 'Neoclassical Revival',
'NS' : 'No Style',
'PR' : 'Prairie',
'QA' : 'Queen Anne',
'SE' : 'Second Empire',
'SCR' : 'Spanish Colonial Revival',
'ST' : 'Stick'
}

def extract_photos_from_gdb(gdb, out_folder, fcname):
    fldBLOB = 'DATA'  
    fldAttName = 'ATT_NAME'
    fldGlobID = 'REL_GLOBALID'    
    #Name of table with attachments
    tbl = os.path.join(gdb, "{}__ATTACH".format(fcname))
    with da.SearchCursor(tbl, [fldBLOB, fldAttName, fldGlobID]) as cursor:
    #iterate through table with attachments
    #Don't sort this cursor...weird things happen
       for row in cursor:
          binaryRep = row[0]
          GlobID = row[2]
          # save to disk
          uniquename = CreateUniqueName(GlobID + ".jpg", out_folder)
          open(uniquename,'wb').write(binaryRep.tobytes())
    
    #Convert all pictures to thumbnails
    os.chdir(out_folder)
    for photo in os.listdir(out_folder):
        fullpath = os.path.abspath(photo)
        fullpathdir = os.path.dirname(fullpath)
        thumbpath = os.path.join(fullpathdir, "thumbnails")
        if not os.path.exists(thumbpath):
            os.makedirs(thumbpath)
        if fullpath.endswith(".jpg"):
            img = Image.open(fullpath)
            img.thumbnail((390,260),Image.ANTIALIAS)
            img.save(os.path.join(thumbpath, os.path.basename(fullpath)), "JPEG")

def getdomaindescription(gdb, subtype, codedvalue):
    #In order to access domains, must make a local copy of database through ArcMap
    dmns = da.ListDomains(gdb)
    for dmn in dmns:
        if dmn.name == subtype:
            vals = dmn.codedValues
            val = vals.get(codedvalue)
            return val

def create_resource_table(gdb, fc, photo_folder, spatref, sql=None):
    #Define table template
    thumbpath = os.path.join(photo_folder, "thumbnails")
    document = docx.Document(os.path.join(os.path.dirname(__file__), 
                                          "../templates/template.docx"))
    infc = os.path.join(gdb, fc)
    flds = ["ResourceID", "PropName", "Address", 
            "SHAPE@X", "SHAPE@Y",  "StrucType", "BldType", 
            "StyleType", "Eligibility", "ConstYr", "Notes", "GlobalID"]
    with da.SearchCursor(infc, flds, sql, spatref) as cursor:
        tbl = document.add_table(rows=1, cols=9)
        hdr_cells  = tbl.rows[0].cells
        hdr_cells[1].text = "Resource ID"
        hdr_cells[2].text = "Name"
        hdr_cells[3].text = "Address"
        hdr_cells[4].text = "UTM E"
        hdr_cells[5].text = "UTM N"
        hdr_cells[6].text = "Type/Style"
        hdr_cells[7].text = "NRHP Evaluation"
        hdr_cells[8].text = "Notes"
        for row in sorted(cursor):
            resid = "{}".format(row[flds.index("ResourceID")])
            resname = "{}".format(row[flds.index("PropName")])
            address = "{}".format(row[flds.index("Address")])
            easting = row[flds.index("SHAPE@X")]
            northing = row[flds.index("SHAPE@Y")]
            structure = row[flds.index("StrucType")]
            bldg = row[flds.index("BldType")]
            stylerow = row[flds.index("StyleType")]
            nrhp = "{}".format(row[flds.index("Eligibility")])
            constyr = "{}".format(row[flds.index("ConstYr")])
            notes = "{}".format(row[flds.index("Notes")])
            glblid = "{}".format(row[flds.index("GlobalID")])
            #add feature count logic
            row_cells = tbl.add_row().cells
            row_cells[1].text = "{}".format(resid)
            #temp picture holder
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run()
            counterbool = True
            piccounter = 0
            pic = os.path.join(thumbpath, "{}.jpg".format(glblid))
            if os.path.exists(pic):
                hgt = 1400000
                run.add_picture(pic, width = hgt * 1.5, height = hgt)
                while counterbool == True:
                    pic = os.path.join(thumbpath, 
                                       "{}{}.jpg".format(glblid, piccounter))
                    if os.path.exists(pic):
                        hgt = 1400000
                        wdth = hgt * 1.5
                        run.add_picture(pic, width=wdth, height=hgt)
                        piccounter += 1
                    else:
                        counterbool = False
            else:
                pass
            try:
                bldgsub = subtypes.get(structure)
                bldgstyle = getdomaindescription(gdb, subtypes.get(structure), bldg)
                styletype = style.get(stylerow)
                eligibility = eligdict.get(nrhp)
                row_cells[2].text = resname
                row_cells[3].text = address
                row_cells[4].text = "{}".format(int(easting))
                row_cells[5].text = "{}".format(int(northing))
                row_cells[6].text = ", ".join((bldgsub, bldgstyle, styletype))
                row_cells[7].text = eligibility
                if not constyr == "None":
                    row_cells[8].text = "; ".join((constyr, '{}'.format(notes)))
                else:
                    row_cells[8].text = "{}".format(notes)
            except ValueError as e:
                print e.message
            document.save(os.path.join(photo_folder,"tblReport.docx"))
        return document

def create_list_from_table_report(doc):
    #Assumes one table in history report
    tbl = doc.tables[0]
    dblist = []
    for row in tbl.rows:
        _toadd = []
        for cell in row.cells:
            _toadd.append(cell.text)
        dblist.append(_toadd)
    return dblist

def spatial_join(geom, feat, fld):
    """
    On the fly spatial join for a given coordinate pair
    """    
    pass

if __name__ == "__main__":
    pass